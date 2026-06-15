"""Build a Word (.docx) version of the combined methodology document."""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

FIG = Path(__file__).resolve().parents[2] / "figures"
OUT = FIG / "methodology_combined.docx"

simple_png  = FIG / "methodology_simple.png"
complex_png = FIG / "methodology.png"

doc = Document()

for section in doc.sections:
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)

def heading(text, size=18, color=(13, 71, 161), bold=True, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(*color)
    return p

def para(text, size=11, italic=False, color=(34, 34, 34), align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    r.font.color.rgb = RGBColor(*color)
    return p

# -------- Title --------
heading("MDIE — Modification-Disentangled Identity Encoder",
        size=22, align=WD_ALIGN_PARAGRAPH.CENTER)
para("- Mrinal", size=12, italic=True, color=(85, 85, 85),
     align=WD_ALIGN_PARAGRAPH.CENTER)

# -------- 1. Simple diagram --------
doc.add_paragraph()
heading("1. Simple methodology — four boxes, mapped to the proposed plan", size=14)
doc.add_picture(str(simple_png), width=Inches(7.0))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# -------- 2. Complex diagram --------
heading("2. Detailed methodology — full training graph", size=14)
doc.add_picture(str(complex_png), width=Inches(7.0))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# -------- 3. Plain-English description --------
heading("3. How the approach works", size=14)

sections = [
    ("The problem",
     "Face-recognition systems trained on clean photos collapse when the same person shows up "
     "with a mask, glasses, post-surgery features, partial occlusion, an aging gap, low-light "
     "capture, or an adversarial perturbation. Stage 1 of the proposed execution plan asks me "
     "to first prove this failure quantitatively. I did, on LFW with 9 modification types: "
     "ArcFace's worst-case AUC drops by 0.198."),
    ("The core idea",
     "I do NOT change inference. The deployed model is still 'embed image then cosine compare', "
     "exactly like ArcFace. All novelty lives at training time, in two extra signals attached "
     "to the same IR-50 backbone."),
    ("Novelty 1 — AMD (Adversarial Modification Disentanglement)",
     "I attach a small classifier head whose job is to guess which modification the input has "
     "(mask? glasses? FGSM?). Between the head and the backbone I insert a Gradient-Reversal "
     "Layer (GRL). The head still tries to predict the modification, but the encoder is pushed "
     "in the OPPOSITE direction — it learns to make the modification cue unreadable. So the "
     "embedding is forced to be modification-agnostic by construction."),
    ("Novelty 2 — ICCL (Identity-Consistency Contrastive Loss)",
     "For every clean image x_i I generate the same identity wearing one of the 9 "
     "modifications, x_tilde_i. ICCL pulls the embeddings z_i and z_tilde_i together while "
     "pushing them apart from other identities. I hard-mine: negatives that share the SAME "
     "modification are weighted 2x, so the model cannot cheat by using the modification "
     "itself as the discriminator."),
    ("Why the two together (and not just one)",
     "AMD removes the modification signal from the features. ICCL guarantees the identity "
     "signal survives that removal. One without the other under-performs in my ablations: "
     "ICCL alone leaks modification info; AMD alone can collapse identity. Together they hit "
     "the sweet spot."),
    ("Algorithms used",
     "Backbone: IR-50 (ResNet-50 variant for face), 512-D embeddings. Identity supervision: "
     "ArcFace additive angular margin (m=0.5, s=64). Novel losses: AMD via Gradient-Reversal "
     "Layer (Ganin and Lempitsky) on a modification-id MLP head; ICCL, an InfoNCE-style "
     "contrastive loss with modification-aware hard-negative mining. Ablation: RATA — "
     "Region-Attention Transformer Aggregator (reported as a negative result in my "
     "small-data regime). Baselines: FaceNet (triplet), ArcFace, CosFace, MobileFaceNet, "
     "and — added in v3 — InsightFace's production w600k_r50 (IR-50 trained on WebFace12M, "
     "auto-downloaded from HuggingFace) as a strong external reference. "
     "Modification engine: surgical nasal warp, surgical jaw warp, opaque glasses, "
     "mouth-and-nose mask, random rectangular occlusion, age-progression filter, "
     "low-light gamma plus blue shift, FGSM adversarial perturbation, plus a clean "
     "control — nine in total. Training stack: PyTorch 2.1+, AMP mixed-precision on an "
     "RTX 3050 4GB, AdamW with cosine schedule and warmup, MTCNN-style 5-point alignment to "
     "a 112x112 crop. Evaluation: cosine-similarity verification, ROC / AUC, EER, and "
     "FAR at fixed FRR; v3 adds Grad-CAM grids and CAM-IoU on the periocular eye region for "
     "interpretability."),
    ("v3 upgrades — real benchmarks, external baseline, interpretability",
     "On top of the original synthetic-LFW protocol I added three things to make the paper "
     "harder for a reviewer to dismiss. (1) Real-data evaluation harness "
     "(src/eval/run_real_benchmarks.py) covering five standard public benchmarks: MFR2 "
     "(real masked faces), CALFW (cross-age, 10-fold), AgeDB-30 (aging, 10-fold), IIITD "
     "Plastic Surgery (gated), and IJB-C occlusion protocol (gated). The free benchmarks "
     "auto-load from InsightFace .bin files dropped into datasets_cache/benchmarks/. "
     "(2) A strong external baseline — InsightFace's production w600k_r50 (IR-50 trained "
     "on WebFace12M, 174 MB) — auto-downloaded from HuggingFace by the eval harness. The "
     "headline comparison in v3 is therefore MDIE vs a production model trained on 12 "
     "million faces, not just my own re-trained baselines. (3) Grad-CAM interpretability "
     "(src/eval/gradcam.py): a grid showing where each model attends on representative "
     "pairs and a CAM-IoU bar chart measuring overlap with the periocular eye-region box. "
     "The hypothesis the figure tests is that MDIE keeps a high CAM-IoU on the eye region "
     "across modifications while ArcFace's attention drifts toward the now-occluded lower "
     "face — the geometric reason masks hurt ArcFace so much."),
    ("Mapping to the proposed execution plan",
     "Stage 1 (problem validation): done — I benchmark FaceNet, ArcFace, CosFace, "
     "MobileFaceNet and report ROC, EER, FAR@FRR on all 9 modifications. "
     "Stage 3A (lightweight edge backbone): IR-50 + ArcFace head, same footprint as a "
     "deployable model. "
     "Stages 3C + 3D (region-stable representation + embedding optimisation): AMD + ICCL "
     "provide exactly this — the encoder learns to attend to identity-stable regions and the "
     "embedding geometry is optimised for cross-modification matching. "
     "Stage 4 (laptop demonstrator): trained and evaluated end-to-end on an RTX 3050 4GB. "
     "Future stages (2 GAN aug, 3B depth+IR, 5–7 edge / federated / field): the architecture "
     "is designed so each plugs in without changing the inference path."),
    ("Headline result",
     "On LFW with 217 identities and 6000 verification pairs per modification: "
     "ArcFace worst-case AUC drop = 0.198. MDIE worst-case AUC drop = 0.027. "
     "That is roughly 7x more robust under degradation, with the same inference cost as "
     "ArcFace. Best per-modification gain: disguise+mask AUC 0.704 → 0.831 (+12.7 pp)."),
    ("How it is publishable",
     "Four things together: (1) a failure-mode benchmark across 9 modifications on standard "
     "data, (2) two genuinely new training-only losses — AMD via GRL on a modification-id "
     "head, and ICCL with modification-aware hard negatives, (3) honest ablations including "
     "a negative result on the RATA region-attention variant in my small-data regime, and "
     "(4) v3 real-benchmark cross-validation against InsightFace's production w600k_r50 "
     "plus Grad-CAM interpretability evidence. The deployed model needs no extra parameters "
     "or labels at inference, which is what makes it directly relevant to the edge-deployment "
     "stages of the plan."),
    ("Honest caveats",
     "RATA — the third originally-proposed component — does not work in my small-data regime "
     "(173 training identities) and is reported as a negative result. It almost certainly "
     "needs MS1MV3-scale training to show its design intent, which does not fit on a 4 GB "
     "card. Pretrained seed: I could not find a public PyTorch IR-50 checkpoint whose key "
     "layout matches face.evoLVe's IR-50, so MDIE is currently trained from random init. For "
     "the eval comparison I use the InsightFace w600k_r50 baseline (a different architecture, "
     "iresnet50) which I do download as a production reference point. A matched-layout "
     "pretrained seed remains future work. The two gated real benchmarks (IIITD Plastic "
     "Surgery, IJB-C) require access agreements and are silently skipped when their "
     "environment variables are unset."),
]

for h, body in sections:
    p = doc.add_paragraph()
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(13, 71, 161)
    para(body, size=10.5)

doc.save(OUT)
print("wrote", OUT)
